#!/usr/bin/env python3
"""
Executive report builder for the website review scaffold.

Reads a JSON file of findings and recommendations and writes a CEO-level Word
report with consistent, board-document formatting (PLAN.md sections 12 and
35): a cover page with a static contents list, a running header with a
different first page, numbered sections, an at-a-glance card strip, the
bottom line as a display-type statement, a scorecard with measured score
bars, hairline-ruled tables with color-chip posture and severity cells, an
evidence appendix, and a numbered footer. Formatting lives here so the
report looks identical on every run regardless of who or what invokes it,
and regardless of which site was analyzed.

Usage:
    python build_exec_report.py <input_json> <output_docx>

Dependency:
    pip install python-docx matplotlib
    (matplotlib is only exercised when the report data carries a quarterly
    trend with three or more quarters of history)
"""

import json
import math
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import report_charts

BODY_FONT = "Calibri"
DISPLAY_FONT = "Georgia"                     # cover title, section numbers, big values
MONO_FONT = "Consolas"

# Single accent color. Swap this hex to rebrand the whole report.
ACCENT_HEX = "0B1F3A"                       # deep navy
ACCENT_RGB = RGBColor(0x0B, 0x1F, 0x3A)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
BODY_RGB = RGBColor(0x26, 0x2B, 0x33)
MUTED_RGB = RGBColor(0x5A, 0x66, 0x72)      # secondary text
BANNER_SUB_RGB = RGBColor(0xB9, 0xC6, 0xDA)  # light blue-gray on navy
MARK_RGB = RGBColor(0xB3, 0x26, 0x1E)       # red, for the marked (problem) text
GOLD_HEX = "C9A227"                          # short cover rule, section numbers
GOLD_RGB = RGBColor(0xC9, 0xA2, 0x27)
HAIRLINE_HEX = "D8DEE9"
CODE_FILL_HEX = "F7F8FA"
CHIP_FILL_HEX = "E6EAF1"                     # neutral chip (effort column)
BAR_EMPTY_HEX = "E1E6EE"                     # unfilled segments of a score bar
SCORE_BAR_SEGMENTS = 12

# Background fill per severity for the severity cell. Muted, print-safe tones.
SEVERITY_FILL = {
    "Critical": "7F1D1D",
    "High": "C0392B",
    "Medium": "E2A800",
    "Low": "1E7B4F",
}
SEVERITY_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}

# Background fill per measured posture band (from the scanner scorecard).
BAND_FILL = {
    "Strong": "1E7B4F",
    "Adequate": "E2A800",
    "Weak": "CB6120",
    "Poor": "B3261E",
    "Not measured": "8A94A6",
}

# Fills light enough to need dark text for contrast; the rest use white.
LIGHT_FILLS = {"E2A800", CHIP_FILL_HEX}

PAGE_WIDTH = Inches(7.0)  # usable width inside 0.75in margins on US Letter


# ---------------------------------------------------------------- low level --

# OOXML property elements require their children in a fixed schema order.
# python-docx reads any order, but Word is strict: appending (say) tcBorders
# after shd makes Word silently repair or mis-lay-out the file. Every manual
# element below is therefore inserted before the tags that must follow it.
_TCPR_AFTER_TCBORDERS = ("w:shd", "w:noWrap", "w:tcMar", "w:textDirection",
                         "w:tcFitText", "w:vAlign", "w:hideMark")
_TCPR_AFTER_SHD = _TCPR_AFTER_TCBORDERS[1:]
_TBLPR_AFTER_TBLBORDERS = ("w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
                           "w:tblCaption", "w:tblDescription")
_TBLPR_AFTER_TBLCELLMAR = ("w:tblLook", "w:tblCaption", "w:tblDescription")
_PPR_AFTER_PBDR = ("w:shd", "w:tabs", "w:suppressAutoHyphens", "w:spacing",
                   "w:ind", "w:contextualSpacing", "w:jc", "w:rPr", "w:sectPr")
_RPR_AFTER_SPACING = ("w:w", "w:kern", "w:position", "w:sz", "w:szCs",
                      "w:highlight", "w:u", "w:vertAlign", "w:lang")


def add_run(paragraph, text, size=10, bold=False, color=BODY_RGB, italic=False,
            font=BODY_FONT, letter_spacing=None):
    # Tolerate a non-string value from hand-authored report data: None is an empty
    # run, any other scalar is stringified rather than crashing the build.
    run = paragraph.add_run("" if text is None else str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = font
    if letter_spacing:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(letter_spacing))  # twentieths of a point
        run._r.get_or_add_rPr().insert_element_before(sp, *_RPR_AFTER_SPACING)
    return run


def shade_cell(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.insert_element_before(shd, *_TCPR_AFTER_SHD)


def set_cell_text(cell, text, size=9, bold=False, color=BODY_RGB, align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if align is not None:
        para.alignment = align
    add_run(para, text, size=size, bold=bold, color=color)
    return para


def set_table_padding(table, top=80, bottom=80, left=110, right=110):
    """Default cell margins for the whole table, in twentieths of a point.
    Roomy cells are the single clearest difference between a designed table
    and Word's cramped default."""
    tbl_pr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.insert_element_before(mar, *_TBLPR_AFTER_TBLCELLMAR)


def set_table_borders(table, **edges):
    """Explicit border per edge (top/left/bottom/right/insideH/insideV).
    Each edge is (size_in_eighth_points, hex_color) or None for no border."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        spec = edges.get(edge)
        if spec:
            size, color = spec
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.insert_element_before(borders, *_TBLPR_AFTER_TBLBORDERS)


def set_cell_border(cell, edge, size, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.insert_element_before(borders, *_TCPR_AFTER_TCBORDERS)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:color"), color)
    borders.append(el)


def keep_rows_together(table, header=False):
    """Stop rows splitting across pages; optionally repeat the header row."""
    for idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        if header and idx == 0:
            tr_pr.append(OxmlElement("w:tblHeader"))


def keep_table_together(table):
    """Move the whole table (and the heading welded to it by keep-with-next)
    to the next page rather than splitting it mid-table. Word satisfies this
    only when the table fits one page, so an oversized table still splits
    instead of overflowing."""
    for row in list(table.rows)[:-1]:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.keep_with_next = True


def set_col_widths(table, widths):
    # python-docx needs the width set on every cell to be reliable.
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def text_color_for_fill(hex_fill):
    return BODY_RGB if hex_fill.upper() in LIGHT_FILLS else WHITE_RGB


def add_page_number_field(paragraph, size=8, color=MUTED_RGB):
    run = paragraph.add_run()
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = BODY_FONT
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


# ------------------------------------------------------------- components --

def section_heading(document, text, number=None):
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(22)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.keep_with_next = True  # never strand a title at a page bottom
    # Thin rule under the heading, drawn as a paragraph bottom border.
    p_pr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), HAIRLINE_HEX)
    borders.append(bottom)
    p_pr.insert_element_before(borders, *_PPR_AFTER_PBDR)
    if number is not None:
        add_run(para, f"{number:02d}", size=13, bold=True, color=GOLD_RGB,
                font=DISPLAY_FONT)
        add_run(para, "   ", size=12)
    add_run(para, text.upper(), size=11.5, bold=True, color=ACCENT_RGB, letter_spacing=20)
    return para


def _rule_paragraph(document, size, color, right_indent=None,
                    space_before=0, space_after=10):
    """A horizontal rule drawn as a paragraph bottom border. A right indent
    shortens the rule (a short accent rule reads as designed; a full-width
    one reads as a divider)."""
    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(space_before)
    rule.paragraph_format.space_after = Pt(space_after)
    if right_indent is not None:
        rule.paragraph_format.right_indent = right_indent
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.insert_element_before(borders, *_PPR_AFTER_PBDR)
    return rule


def _scope_text(scope):
    bits = []
    pages = (scope or {}).get("pages_reviewed")
    if pages:
        bits.append(f"{pages} page(s) reviewed")
    if (scope or {}).get("method"):
        bits.append(str(scope["method"]))  # coerce so a non-string method cannot crash join
    return "  |  ".join(bits)


def add_cover(document, data, section_titles):
    """A full cover page: whitespace, kicker, the site name at display scale,
    a short gold rule, the measured overall posture, meta lines, and a static
    contents list of the sections that actually render. No Word TOC fields:
    they show empty until manually refreshed and read as broken."""
    site = data.get("site", "Website")

    breather = document.add_paragraph()
    breather.paragraph_format.space_before = Pt(96)
    breather.paragraph_format.space_after = Pt(0)

    # An optional report_label ("SAMPLE REPORT") replaces the kicker's second
    # half and gets an unmissable gold badge, so a demonstration copy can
    # never be mistaken for a client deliverable.
    label = (str(data.get("report_label") or "") or "EXECUTIVE REPORT").upper()
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(10)
    add_run(kicker, f"WEBSITE REVIEW  /  {label}", size=9,
            bold=True, color=MUTED_RGB, letter_spacing=44)
    if data.get("report_label"):
        badge = document.add_paragraph()
        badge.paragraph_format.space_after = Pt(12)
        chip = add_run(badge, f"  {label}  ", size=11, bold=True,
                       color=ACCENT_RGB, letter_spacing=30)
        _shade_run(chip, GOLD_HEX)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    add_run(title, site, size=38, bold=True, color=ACCENT_RGB, font=DISPLAY_FONT)

    _rule_paragraph(document, size=28, color=GOLD_HEX,
                    right_indent=PAGE_WIDTH - Inches(2.3), space_after=18)

    overall = (data.get("scorecard") or {}).get("overall")
    if overall:
        overall = str(overall)  # tolerate a non-string band from hand-authored data
        posture = document.add_paragraph()
        posture.paragraph_format.space_after = Pt(16)
        add_run(posture, "OVERALL POSTURE   ", size=9, bold=True,
                color=MUTED_RGB, letter_spacing=24)
        fill = BAND_FILL.get(_hkey(overall), BAND_FILL["Not measured"])
        chip = add_run(posture, f"  {overall.upper()}  ", size=9, bold=True,
                       color=text_color_for_fill(fill), letter_spacing=24)
        _shade_run(chip, fill)

    meta_lines = [data.get("target_url", ""), data.get("date", ""),
                  _scope_text(data.get("scope"))]
    for line in meta_lines:
        if not line:
            continue
        meta = document.add_paragraph()
        meta.paragraph_format.space_after = Pt(2)
        add_run(meta, line, size=10, color=MUTED_RGB)

    if section_titles:
        gap = document.add_paragraph()
        gap.paragraph_format.space_before = Pt(52)
        gap.paragraph_format.space_after = Pt(0)
        label = document.add_paragraph()
        label.paragraph_format.space_after = Pt(8)
        add_run(label, "IN THIS REPORT", size=8.5, bold=True,
                color=MUTED_RGB, letter_spacing=30)
        for n, name in enumerate(section_titles, start=1):
            item = document.add_paragraph()
            item.paragraph_format.space_after = Pt(4)
            add_run(item, f"{n:02d}", size=10, bold=True, color=GOLD_RGB,
                    font=DISPLAY_FONT)
            add_run(item, f"   {name}", size=10.5, color=BODY_RGB)

    closing = document.add_paragraph()
    closing.paragraph_format.space_before = Pt(40)
    closing.paragraph_format.space_after = Pt(0)
    # A sample with illustrative history must not carry the measured-only
    # pledge; cover_note lets the data state what is true for this copy.
    note = data.get("cover_note") or (
        "Prepared for executive review. Every finding in this "
        "document cites a measured check; nothing is estimated.")
    add_run(closing, note, size=8.5, italic=True, color=MUTED_RGB)

    document.add_page_break()


def add_running_header(document, site, date, label=None):
    """A small right-aligned running header over a hairline on content pages.
    The cover page (different first page) carries nothing. A report_label
    replaces the default "Website Review" so every page names the copy."""
    section = document.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0]
    para.text = ""
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), HAIRLINE_HEX)
    borders.append(bottom)
    p_pr.insert_element_before(borders, *_PPR_AFTER_PBDR)
    bits = [site, label or "Website Review"] + ([date] if date else [])
    add_run(para, "   |   ".join(str(b) for b in bits), size=8, color=MUTED_RGB)
    # Content page numbering starts at 1 (the cover is page 0, unnumbered).
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "0")
    section._sectPr.insert_element_before(pg, "w:cols", "w:docGrid")


def add_progress_strip(document, progress):
    """The re-review headline: what changed since the previous report."""
    prev = progress.get("previous_date")
    resolved = progress.get("resolved_issues")
    new = progress.get("new_issues")
    if resolved is None and new is None:
        return
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(0)
    add_run(para, "Since the previous review" + (f" ({prev})" if prev else "") + ":  ",
            size=9.5, color=MUTED_RGB)
    add_run(para, f"{resolved or 0} resolved", size=9.5, bold=True,
            color=RGBColor.from_string(SEVERITY_FILL["Low"]))
    add_run(para, "   ", size=9.5)
    add_run(para, f"{new or 0} new", size=9.5, bold=True,
            color=RGBColor.from_string(SEVERITY_FILL["High"] if new else "5A6672"))


DIRECTION_STYLE = {"improved": ("STRONGER", SEVERITY_FILL["Low"]),
                   "held": ("HELD", "8A94A6"),
                   "declined": ("WEAKER", SEVERITY_FILL["High"])}


def _fmt_score(v):
    return f"{v:.2f}" if isinstance(v, (int, float)) else "n/a"


def add_trend_table(document, delta):
    """Quarter-over-quarter posture: prior band, current band, score
    movement, and a direction chip per category."""
    def trend_row(row):
        def write(cells):
            set_cell_text(cells[0], row.get("category", ""), size=9, bold=True)
            for cell, band in ((cells[1], row.get("prev_band")),
                               (cells[2], row.get("band"))):
                band = band or "Not measured"
                _chip(cell, band, BAND_FILL.get(_hkey(band), BAND_FILL["Not measured"]))
            p, c = row.get("prev_score"), row.get("score")
            change = (f"{c - p:+.2f}"
                      if isinstance(p, (int, float)) and isinstance(c, (int, float))
                      else "n/a")
            set_cell_text(cells[3], f"{_fmt_score(p)} to {_fmt_score(c)} ({change})",
                          size=8.5, color=MUTED_RGB)
            label, fill = DIRECTION_STYLE.get(_hkey(row.get("direction")),
                                              DIRECTION_STYLE["held"])
            _chip(cells[4], label, fill)
        return write

    headers = ["Area", delta.get("prev_quarter") or "Prior",
               delta.get("quarter") or "Current", "Score", "Direction"]
    add_data_table(document, headers,
                   [Inches(1.35), Inches(1.1), Inches(1.1),
                    Inches(2.15), Inches(1.3)],
                   [trend_row(r) for r in delta.get("scorecard", [])])


def add_trend_section(document, trend, chart_dir, prefix, number):
    """Progress this quarter: the QoQ posture table, trend charts (three or
    more quarterly points), and every resolved finding named in full. This
    is the retainer's value story, so it leads the report."""
    section_heading(document, "Progress this quarter", number)
    delta = trend.get("latest_delta") or {}
    if delta.get("scorecard"):
        add_trend_table(document, delta)
        ps = delta.get("pages_scanned") or {}
        if ps.get("prev") is not None or ps.get("current") is not None:
            note = document.add_paragraph()
            note.paragraph_format.space_before = Pt(4)
            add_run(note, f"Pages reviewed: {ps.get('prev')} in "
                          f"{delta.get('prev_quarter')}, {ps.get('current')} "
                          f"in {delta.get('quarter')}.",
                    size=8, color=MUTED_RGB, italic=True)

    # Charts start at three quarterly points; a two-point line implies a
    # slope one interval cannot support, so the table carries the QoQ story.
    if len(trend.get("quarters") or []) >= 3:
        if not prefix:
            raise ValueError("data['slug'] is required to name trend chart "
                             "files; refusing to guess a shared prefix that "
                             "could collide across clients")
        charts = report_charts.render_trend_charts(trend, chart_dir, prefix)
        for chart in charts:
            cap = document.add_paragraph()
            cap.paragraph_format.space_before = Pt(10)
            cap.paragraph_format.space_after = Pt(4)
            cap.paragraph_format.keep_with_next = True
            add_run(cap, chart["caption"], size=8.5, color=MUTED_RGB, italic=True)
            add_framed_picture(document, chart["path"])

    movement = document.add_paragraph()
    movement.paragraph_format.space_before = Pt(10)
    add_run(movement,
            f"{delta.get('resolved_findings', 0)} finding(s) resolved since "
            f"{delta.get('prev_quarter', 'the previous quarter')}; "
            f"{delta.get('new_findings', 0)} new.", size=10, bold=True)
    for item in delta.get("resolved_examples") or []:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        add_run(p, item, size=9)


def _as_rows(items, text_key):
    """Coerce a list of report items to dicts so a hand-authored list of plain
    strings (recommendations/findings/action_plan are human-edited on top of the
    draft, and the sibling quick_wins IS a string list) still renders instead of
    crashing .get() on a non-dict. A bare string becomes {text_key: s} so its text
    appears; a dict passes through; any other scalar becomes an empty row. A WHOLE
    field given as a single value (not a list) is treated as ONE item, whatever its
    type: a lone dict (one finding authored without the wrapping list) renders as
    one row preserving its content, never iterated over its keys; a lone string is
    not iterated per character; a lone scalar does not crash the build."""
    if items is None:
        return []
    if not isinstance(items, (list, tuple)):
        items = [items]
    rows = []
    for it in items:
        if isinstance(it, dict):
            rows.append(it)
        elif isinstance(it, str):
            rows.append({text_key: it})
        else:
            rows.append({})
    return rows


def _as_str_list(value):
    """Coerce a field meant to be a list of strings (quick_wins, assessment
    strengths/weaknesses). A bare string becomes a one-item list, never iterated
    per character; None or an empty/whitespace string becomes []; a list or tuple
    is returned as a list (item text coercion is add_run's job)."""
    if value is None:
        return []
    if isinstance(value, str):
        return [value] if value.strip() else []
    if isinstance(value, (list, tuple)):
        return list(value)
    return [value]


def _finite_number(x):
    """True only for a real, finite int/float. Excludes NaN and +/-Infinity,
    which json.loads accepts by default and which crash round() in the score bar."""
    return isinstance(x, (int, float)) and math.isfinite(x)


def _hkey(x):
    """A band/severity/rating value used as a color-lookup key. A non-string (a
    hand-authored list/dict where a label was expected) is stringified so the
    lookup falls through to its default instead of raising TypeError: unhashable."""
    return x if isinstance(x, str) else str(x)


def _severity_breakdown(findings):
    counts = {}
    for f in findings:
        sev = _hkey(f.get("severity", "Low"))
        counts[sev] = counts.get(sev, 0) + 1
    ordered = sorted(counts.items(), key=lambda kv: SEVERITY_ORDER.get(kv[0], 9))
    return " / ".join(f"{n} {sev}" for sev, n in ordered)


def add_glance_tiles(document, data):
    """One row of at-a-glance tiles, every value copied or counted from the
    supplied data. Nothing here is invented."""
    findings = data.get("findings", [])
    recs = data.get("recommendations", [])
    scorecard = data.get("scorecard") or {}
    rows = scorecard.get("rows", [])
    overall = str(scorecard.get("overall") or "Not measured")  # tolerate a non-string band

    scope = data.get("scope") or {}
    if scope.get("pages_reviewed"):
        fourth = ("PAGES REVIEWED", str(scope["pages_reviewed"]),
                  "in-scope pages scanned", ACCENT_RGB)
    else:
        fourth = ("AREAS MEASURED", str(len(rows)),
                  "scan categories" if rows else "no scorecard supplied", ACCENT_RGB)
    tiles = [
        ("OVERALL POSTURE", overall, "measured scan verdicts",
         RGBColor.from_string(BAND_FILL.get(_hkey(overall), BAND_FILL["Not measured"]))),
        ("KEY FINDINGS", str(len(findings)),
         _severity_breakdown(findings) or "none recorded", ACCENT_RGB),
        ("RECOMMENDATIONS", str(len(recs)),
         "ranked by priority" if recs else "to be authored", ACCENT_RGB),
        fourth,
    ]

    table = document.add_table(rows=1, cols=len(tiles))
    set_table_borders(table, bottom=(4, HAIRLINE_HEX), left=(4, HAIRLINE_HEX),
                      right=(4, HAIRLINE_HEX), insideV=(4, HAIRLINE_HEX))
    set_table_padding(table, top=110, bottom=110, left=130, right=100)
    for cell, (label, value, sub, value_color) in zip(table.rows[0].cells, tiles):
        set_cell_border(cell, "top", 20, ACCENT_HEX)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(3)
        add_run(p1, label, size=7.5, bold=True, color=MUTED_RGB, letter_spacing=16)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        add_run(p2, value, size=15, bold=True, color=value_color, font=DISPLAY_FONT)
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        add_run(p3, sub, size=7.5, color=MUTED_RGB)
    set_col_widths(table, [Inches(1.75)] * len(tiles))
    keep_rows_together(table)


VITALS_FILL = {"Good": BAND_FILL["Strong"], "Needs work": BAND_FILL["Adequate"],
               "Poor": BAND_FILL["Poor"]}


def add_vitals_panel(document, web_vitals):
    """A compact metric strip for Core Web Vitals: one bordered tile per
    metric with the value large, the label small, and a rating chip."""
    metrics = web_vitals.get("metrics") or []
    if not metrics:
        return
    table = document.add_table(rows=1, cols=len(metrics))
    set_table_borders(table, top=(4, HAIRLINE_HEX), bottom=(4, HAIRLINE_HEX),
                      left=(4, HAIRLINE_HEX), right=(4, HAIRLINE_HEX),
                      insideV=(4, HAIRLINE_HEX))
    set_table_padding(table, top=110, bottom=110, left=130, right=130)
    for cell, m in zip(table.rows[0].cells, metrics):
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        add_run(p1, m.get("label", ""), size=8, bold=True, color=MUTED_RGB, letter_spacing=12)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        add_run(p2, str(m.get("value", "")), size=17, bold=True, color=ACCENT_RGB,
                font=DISPLAY_FONT)
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        rating = m.get("rating", "")
        fill = VITALS_FILL.get(_hkey(rating), BAND_FILL["Not measured"])
        run = add_run(p3, f" {rating} ", size=7.5, bold=True, color=text_color_for_fill(fill))
        _shade_run(run, fill)
    set_col_widths(table, [Inches(7.0 / len(metrics))] * len(metrics))
    keep_rows_together(table)
    src = document.add_paragraph()
    src.paragraph_format.space_before = Pt(3)
    src.paragraph_format.space_after = Pt(0)
    add_run(src, web_vitals.get("captured_note", ""), size=8, italic=True, color=MUTED_RGB)


def add_key_dates_panel(document, key_dates):
    """A card strip of conversation-starter dates (certificate and domain
    renewal, domain age). Same white-card language as the vitals panel; the
    value is a date, with a small relative-time line beneath it."""
    items = key_dates.get("items") or []
    if not items:
        return
    table = document.add_table(rows=1, cols=len(items))
    set_table_borders(table, top=(4, HAIRLINE_HEX), bottom=(4, HAIRLINE_HEX),
                      left=(4, HAIRLINE_HEX), right=(4, HAIRLINE_HEX),
                      insideV=(4, HAIRLINE_HEX))
    set_table_padding(table, top=110, bottom=110, left=130, right=130)
    for cell, item in zip(table.rows[0].cells, items):
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        add_run(p1, str(item.get("label") or "").upper(), size=8, bold=True,
                color=MUTED_RGB, letter_spacing=12)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        add_run(p2, str(item.get("value", "")), size=14, bold=True,
                color=ACCENT_RGB, font=DISPLAY_FONT)
        detail = item.get("detail", "")
        if detail:
            p3 = cell.add_paragraph()
            p3.paragraph_format.space_after = Pt(0)
            add_run(p3, detail, size=8, color=MUTED_RGB)
    set_col_widths(table, [Inches(7.0 / len(items))] * len(items))
    keep_rows_together(table)
    note = key_dates.get("note")
    if note:
        src = document.add_paragraph()
        src.paragraph_format.space_before = Pt(3)
        src.paragraph_format.space_after = Pt(0)
        add_run(src, note, size=8, italic=True, color=MUTED_RGB)


def _shade_run(run, hex_fill):
    """Highlight one run with a solid fill (a chip inside a paragraph)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    run._r.get_or_add_rPr().append(shd)


def add_callout(document, text):
    """The bottom line as a statement: display type on white behind a heavy
    navy left border, not a filled box. The conclusion should read like a
    sentence the CEO can quote, not a UI component."""
    table = document.add_table(rows=1, cols=1)
    set_table_borders(table)
    set_table_padding(table, top=60, bottom=60, left=220, right=120)
    cell = table.rows[0].cells[0]
    set_cell_border(cell, "left", 30, ACCENT_HEX)
    cell.text = ""
    kicker = cell.paragraphs[0]
    kicker.paragraph_format.space_after = Pt(5)
    add_run(kicker, "THE BOTTOM LINE", size=8, bold=True, color=MUTED_RGB,
            letter_spacing=30)
    stmt = cell.add_paragraph()
    stmt.paragraph_format.space_after = Pt(0)
    stmt.paragraph_format.line_spacing = 1.3
    add_run(stmt, text, size=13, color=BODY_RGB, font=DISPLAY_FONT)
    set_col_widths(table, [PAGE_WIDTH])


def add_assessment(document, assessment):
    """Two columns: measured strengths on the left, priorities to fix on the
    right, each a short bulleted list."""
    strengths = _as_str_list(assessment.get("strengths"))
    weaknesses = _as_str_list(assessment.get("weaknesses"))
    if not strengths and not weaknesses:
        return
    table = document.add_table(rows=2, cols=2)
    set_table_borders(table, insideV=(4, HAIRLINE_HEX))
    set_table_padding(table, top=60, bottom=90, left=130, right=130)
    for cell, title, hex_color in ((table.rows[0].cells[0], "STRENGTHS", BAND_FILL["Strong"]),
                                   (table.rows[0].cells[1], "PRIORITIES TO FIX", BAND_FILL["Poor"])):
        para = set_cell_text(cell, title, size=8.5, bold=True,
                             color=RGBColor.from_string(hex_color))
        para.runs[0].font.name = BODY_FONT
        set_cell_border(cell, "bottom", 10, hex_color)
    for cell, items in ((table.rows[1].cells[0], strengths),
                        (table.rows[1].cells[1], weaknesses)):
        cell.text = ""
        for idx, item in enumerate(items or ["None recorded."]):
            para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            para.paragraph_format.space_after = Pt(3)
            add_run(para, "- ", size=9, bold=True, color=ACCENT_RGB)
            add_run(para, item, size=9)
    set_col_widths(table, [Inches(3.5), Inches(3.5)])
    keep_rows_together(table)


def add_data_table(document, headers, widths, rows):
    """A hairline-ruled table: navy header row, horizontal rules only, roomy
    cells. `rows` is a list of cell writers, one callable per row taking the
    row's cells."""
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, bottom=(6, HAIRLINE_HEX), insideH=(4, HAIRLINE_HEX))
    set_table_padding(table)
    for cell, label in zip(table.rows[0].cells, headers):
        shade_cell(cell, ACCENT_HEX)
        set_cell_text(cell, label.upper(), size=8, bold=True, color=WHITE_RGB)
    for write_row in rows:
        write_row(table.add_row().cells)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_col_widths(table, widths)
    keep_rows_together(table, header=True)
    return table


def _chip(cell, text, fill):
    shade_cell(cell, fill)
    set_cell_text(cell, text, size=8.5, bold=True, color=text_color_for_fill(fill),
                  align=WD_ALIGN_PARAGRAPH.CENTER)


SCORE_BAR_BLOCK = chr(0x2588)  # FULL BLOCK; solid segments print cleanly


def set_score_bar_cell(cell, score, band):
    """A measured score as a segmented bar in the band's color. Drawn only
    from the numeric score the scan produced; never derived from the band."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    filled = max(0, min(SCORE_BAR_SEGMENTS, round(score * SCORE_BAR_SEGMENTS)))
    band_hex = BAND_FILL.get(_hkey(band), BAND_FILL["Not measured"])
    if filled:
        add_run(para, SCORE_BAR_BLOCK * filled, size=8,
                color=RGBColor.from_string(band_hex), font=MONO_FONT)
    if filled < SCORE_BAR_SEGMENTS:
        add_run(para, SCORE_BAR_BLOCK * (SCORE_BAR_SEGMENTS - filled), size=8,
                color=RGBColor.from_string(BAR_EMPTY_HEX), font=MONO_FONT)
    add_run(para, f"  {score:.2f}", size=8, color=MUTED_RGB)


def _add_marked_runs(paragraph, text, highlights, size=8):
    """Add `text` as monospace runs, marking every occurrence of any highlight
    substring with a yellow highlighter and bold red text so the exact problem
    is unmissable."""
    def style(run, marked=False):
        run.font.name = MONO_FONT
        run.font.size = Pt(size)
        if marked:
            run.font.bold = True
            run.font.color.rgb = MARK_RGB
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            run.font.color.rgb = BODY_RGB

    highlights = [h for h in (highlights or []) if h]
    i, n = 0, len(text)
    while i < n:
        nxt, nxt_h = None, None
        for h in highlights:
            idx = text.find(h, i)
            if idx != -1 and (nxt is None or idx < nxt):
                nxt, nxt_h = idx, h
        if nxt is None:
            style(paragraph.add_run(text[i:]))
            break
        if nxt > i:
            style(paragraph.add_run(text[i:nxt]))
        style(paragraph.add_run(nxt_h), marked=True)
        i = nxt + len(nxt_h)


def add_code_block(document, code, highlight=None):
    """A shaded monospace box holding a literal snippet, with the exact problem
    substring(s) highlighted. `highlight` is a string or list of strings. Each
    source line becomes its own paragraph so line breaks are preserved."""
    # Normalize to a list of strings so a hand-authored numeric highlight (or a list
    # with a numeric item) does not crash text.find/iteration in _add_marked_runs.
    if highlight is None:
        highlight = []
    elif isinstance(highlight, str):
        highlight = [highlight]
    elif isinstance(highlight, (list, tuple)):
        highlight = [str(h) for h in highlight]
    else:
        highlight = [str(highlight)]
    table = document.add_table(rows=1, cols=1)
    set_table_borders(table, top=(4, HAIRLINE_HEX), bottom=(4, HAIRLINE_HEX),
                      left=(4, HAIRLINE_HEX), right=(4, HAIRLINE_HEX))
    set_table_padding(table, top=110, bottom=110, left=150, right=150)
    cell = table.rows[0].cells[0]
    shade_cell(cell, CODE_FILL_HEX)
    cell.text = ""
    for idx, line in enumerate(str(code).split("\n")):  # tolerate a non-string snippet
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        _add_marked_runs(para, line, highlight)
    set_col_widths(table, [PAGE_WIDTH])


def add_framed_picture(document, image_path):
    """A screenshot inside a hairline-framed cell so it reads as an exhibit
    instead of bleeding into the page."""
    table = document.add_table(rows=1, cols=1)
    set_table_borders(table, top=(4, HAIRLINE_HEX), bottom=(4, HAIRLINE_HEX),
                      left=(4, HAIRLINE_HEX), right=(4, HAIRLINE_HEX))
    set_table_padding(table, top=40, bottom=40, left=40, right=40)
    cell = table.rows[0].cells[0]
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    run.add_picture(image_path, width=PAGE_WIDTH - Inches(0.12))
    set_col_widths(table, [PAGE_WIDTH])


def add_footer(document, site):
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.text = ""
    p_pr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:color"), HAIRLINE_HEX)
    borders.append(top)
    p_pr.insert_element_before(borders, *_PPR_AFTER_PBDR)
    para.paragraph_format.tab_stops.add_tab_stop(PAGE_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    add_run(para, f"{site} Website Review", size=8, color=MUTED_RGB)
    add_run(para, "\tPage ", size=8, color=MUTED_RGB)
    add_page_number_field(para)


# ------------------------------------------------------------------ build --

def build(data, out_path, chart_dir=None):
    document = Document()
    chart_dir = Path(chart_dir) if chart_dir else Path(out_path).parent

    # Normalize the human-authored list fields once, up front, so every consumer
    # (the glance tiles' severity breakdown, the section-title logic, and the
    # tables) sees dicts and a hand-authored string list never crashes the build.
    data["findings"] = _as_rows(data.get("findings"), "finding")
    data["recommendations"] = _as_rows(data.get("recommendations"), "recommendation")
    data["action_plan"] = _as_rows(data.get("action_plan"), "action")
    # evidence is authored entirely by hand (draft() never emits it), so it is the
    # most likely to arrive as a list of strings/paths; normalize it and the nested
    # panel lists the same way so a bare-string item renders or degrades, never .get()s.
    data["evidence"] = _as_rows(data.get("evidence"), "caption")
    if isinstance(data.get("web_vitals"), dict):
        data["web_vitals"]["metrics"] = _as_rows(data["web_vitals"].get("metrics"), "label")
    if isinstance(data.get("key_dates"), dict):
        data["key_dates"]["items"] = _as_rows(data["key_dates"].get("items"), "label")
    if isinstance(data.get("scorecard"), dict):
        data["scorecard"]["rows"] = _as_rows(data["scorecard"].get("rows"), "category")
    data["quick_wins"] = _as_str_list(data.get("quick_wins"))
    # Container fields a section reads with .get(): a non-dict (a list or scalar
    # from a hand-authored shape slip) would AttributeError mid-build. Coerce to
    # {} so the section is skipped cleanly, mirroring the list normalization above.
    for _key in ("assessment", "scorecard", "web_vitals", "key_dates", "progress", "scope"):
        if _key in data and not isinstance(data[_key], dict):
            data[_key] = {}

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = BODY_RGB

    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    site = data.get("site", "Website")
    document.core_properties.title = f"{site} Website Review"

    # The section list drives both the cover contents and the heading numbers,
    # so the cover never promises a section that does not render.
    bottom_line = data.get("bottom_line", "")
    assessment = data.get("assessment")
    scorecard = data.get("scorecard")
    web_vitals = data.get("web_vitals")
    progress = data.get("progress") or {}
    trend = progress.get("trend")
    if not isinstance(trend, dict):
        # A hand-authored non-dict trend (string/list/scalar) would AttributeError
        # inside add_trend_section; treat it as no trend so the section is skipped,
        # exactly as the S4 coercion does for the top-level containers. This also
        # lets a real progress strip render (has_exec_summary sees trend as absent).
        trend = None
    # The progress strip (progress without a full quarterly trend chart) belongs to
    # the Executive summary, so that section exists whenever the strip renders -
    # otherwise progress-only data floats a strip under the glance tiles with no
    # heading or contents entry.
    has_exec_summary = bool(bottom_line or assessment or (progress and not trend))
    section_titles = []
    if has_exec_summary:
        section_titles.append("Executive summary")
    if trend:
        section_titles.append("Progress this quarter")
    if scorecard and scorecard.get("rows"):
        section_titles.append("Measured posture")
    if web_vitals and web_vitals.get("metrics"):
        section_titles.append("Core Web Vitals")
    key_dates = data.get("key_dates")
    if key_dates and key_dates.get("items"):
        section_titles.append("Key dates")
    if data.get("findings"):
        section_titles.append("Key findings hurting the site")
    if data.get("recommendations"):
        section_titles.append("Preferred recommendations")
    elif data.get("action_plan"):
        section_titles.append("Recommended plan of action")
    if data.get("quick_wins"):
        section_titles.append("Quick wins")
    if data.get("evidence"):
        section_titles.append("Evidence appendix")
    number_of = {title: n for n, title in enumerate(section_titles, start=1)}

    add_cover(document, data, section_titles)
    add_running_header(document, site, data.get("date", ""),
                       label=data.get("report_label"))
    add_footer(document, site)
    add_glance_tiles(document, data)

    # Executive summary: bottom line, progress, then strengths vs priorities
    if has_exec_summary:
        section_heading(document, "Executive summary",
                        number_of.get("Executive summary"))
    if bottom_line:
        add_callout(document, bottom_line)
    if progress and not trend:
        add_progress_strip(document, progress)
    if assessment:
        add_assessment(document, assessment)

    if trend:
        add_trend_section(document, trend, chart_dir, data.get("slug"),
                          number_of.get("Progress this quarter"))

    # Measured posture scorecard (optional; only when the scan supplied one)
    if scorecard and scorecard.get("rows"):
        heading = "Measured posture"
        overall = scorecard.get("overall")
        if overall:
            heading += f"  (overall: {overall})"
        section_heading(document, heading, number_of.get("Measured posture"))

        rows = scorecard["rows"]
        has_scores = any(_finite_number(r.get("score")) for r in rows)

        def scorecard_row(row):
            def write(cells):
                set_cell_text(cells[0], row.get("category", ""), size=9, bold=True)
                band = row.get("band", "Not measured")
                _chip(cells[1], band, BAND_FILL.get(_hkey(band), BAND_FILL["Not measured"]))
                detail = str(row.get("detail") or "")
                if has_scores:
                    score = row.get("score")
                    if _finite_number(score):
                        set_score_bar_cell(cells[2], score, band)
                        # The bar already shows the score; drop the redundant
                        # "(score N)" suffix from the detail string.
                        detail = re.sub(r"\s*\(score [^)]*\)\s*$", "", detail)
                    else:
                        set_cell_text(cells[2], "not measured", size=8, color=MUTED_RGB)
                set_cell_text(cells[-1], detail, size=8.5, color=MUTED_RGB)
            return write

        if has_scores:
            headers = ["Area", "Posture", "Measured score", "Detail"]
            widths = [Inches(1.4), Inches(1.0), Inches(1.7), Inches(2.9)]
        else:
            headers = ["Area", "Posture", "Measured detail"]
            widths = [Inches(1.75), Inches(1.1), Inches(4.15)]
        add_data_table(document, headers, widths,
                       [scorecard_row(r) for r in rows])

    # Core Web Vitals (real-user or lab), when measured
    if web_vitals and web_vitals.get("metrics"):
        section_heading(document, "Core Web Vitals", number_of.get("Core Web Vitals"))
        add_vitals_panel(document, web_vitals)

    # Key dates: conversation-starter facts (certificate and domain renewal)
    if key_dates and key_dates.get("items"):
        section_heading(document, "Key dates", number_of.get("Key dates"))
        add_key_dates_panel(document, key_dates)

    # Key findings, most severe first
    findings = sorted(
        data.get("findings", []),
        key=lambda f: SEVERITY_ORDER.get(_hkey(f.get("severity", "Low")), 9),
    )
    if findings:
        section_heading(document, "Key findings hurting the site",
                        number_of.get("Key findings hurting the site"))

        def finding_row(row):
            def write(cells):
                sev = row.get("severity", "Low")
                _chip(cells[0], sev, SEVERITY_FILL.get(_hkey(sev), SEVERITY_FILL["Low"]))
                set_cell_text(cells[1], row.get("area", ""), size=9, bold=True)
                set_cell_text(cells[2], row.get("finding", ""), size=9)
                set_cell_text(cells[3], row.get("evidence", ""), size=7.5, color=MUTED_RGB)
            return write

        add_data_table(document, ["Severity", "Area", "Finding", "Evidence"],
                       [Inches(0.85), Inches(1.0), Inches(3.6), Inches(1.55)],
                       [finding_row(f) for f in findings])

    # Preferred recommendations, by rank. Coerce to a numeric key so a hand-authored
    # string rank ("1", "2") mixed with an int or a missing rank never makes sorted()
    # compare str and int and crash the build; a non-numeric or absent rank sorts last.
    def _rank_key(r):
        try:
            return (0, float(r.get("rank")))
        except (TypeError, ValueError):
            return (1, 0.0)
    recs = sorted(data.get("recommendations", []), key=_rank_key)
    if recs:
        section_heading(document, "Preferred recommendations",
                        number_of.get("Preferred recommendations"))

        def rec_row(row):
            def write(cells):
                set_cell_text(cells[0], str(row.get("rank", "")), size=10, bold=True,
                              color=ACCENT_RGB, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(cells[1], row.get("recommendation", ""), size=9)
                set_cell_text(cells[2], row.get("impact", ""), size=8.5, color=MUTED_RGB)
                _chip(cells[3], str(row.get("effort", "")), CHIP_FILL_HEX)
            return write

        table = add_data_table(document, ["#", "Recommendation", "Expected impact", "Effort"],
                               [Inches(0.4), Inches(3.35), Inches(2.5), Inches(0.75)],
                               [rec_row(r) for r in recs])
        # The ranked plan reads as one unit; start it on a fresh page rather
        # than splitting it when it does not fit the space left.
        keep_table_together(table)

    # Plan of action: the measured-derived plan, shown when no hand-authored
    # recommendations exist so a raw run still ships a prioritized plan.
    action_plan = data.get("action_plan")
    if action_plan and not recs:
        section_heading(document, "Recommended plan of action",
                        number_of.get("Recommended plan of action"))

        def action_row(n, row):
            def write(cells):
                set_cell_text(cells[0], str(n), size=10, bold=True,
                              color=ACCENT_RGB, align=WD_ALIGN_PARAGRAPH.CENTER)
                pr = row.get("priority", "Medium")
                _chip(cells[1], pr, SEVERITY_FILL.get(_hkey(pr), SEVERITY_FILL["Low"]))
                set_cell_text(cells[2], row.get("action", ""), size=9)
                set_cell_text(cells[3], row.get("affects", ""), size=8.5, color=MUTED_RGB)
            return write

        table = add_data_table(document, ["#", "Priority", "Action", "Affects"],
                               [Inches(0.4), Inches(0.9), Inches(4.15), Inches(1.05)],
                               [action_row(n, r) for n, r in enumerate(action_plan, start=1)])
        keep_table_together(table)

    # Quick wins
    quick = data.get("quick_wins", [])
    if quick:
        section_heading(document, "Quick wins", number_of.get("Quick wins"))
        for item in quick:
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_run(p, item, size=10)

    # Evidence appendix (optional): captioned proof for significant findings.
    # Each item is {"caption": "..."} plus either an "image" (screenshot path)
    # or a "code" snippet with an optional "highlight" (string or list) that
    # marks the exact problem so the reader sees precisely where the error is.
    evidence = data.get("evidence", [])
    if evidence:
        document.add_page_break()
        section_heading(document, "Evidence appendix",
                        number_of.get("Evidence appendix"))
        for number, item in enumerate(evidence, start=1):
            cap = document.add_paragraph()
            cap.paragraph_format.space_before = Pt(12)
            cap.paragraph_format.space_after = Pt(4)
            cap.paragraph_format.keep_with_next = True  # caption stays with its exhibit
            add_run(cap, f"Exhibit {number}.  ", size=9.5, bold=True, color=ACCENT_RGB)
            add_run(cap, item.get("caption", ""), size=9.5, bold=True)
            if item.get("code") is not None:
                add_code_block(document, item["code"], item.get("highlight"))
            elif isinstance(item.get("image"), str):  # a non-string path is skipped, not Path()'d
                img = item["image"]
                if Path(img).exists():
                    try:
                        add_framed_picture(document, img)
                    except Exception as e:
                        add_run(document.add_paragraph(),
                                f"[image not embedded: {e}]", size=8, italic=True)
                else:
                    add_run(document.add_paragraph(),
                            f"[missing image: {img}]", size=8, italic=True)

    document.save(str(out_path))


def main():
    if len(sys.argv) != 3:
        print("Usage: python build_exec_report.py <input_json> <output_docx>")
        sys.exit(1)
    in_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    if not in_path.is_file():
        print(f"Input JSON not found: {in_path}")
        sys.exit(1)
    try:
        data = json.loads(in_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        print(f"Invalid JSON in {in_path}: {e}")
        sys.exit(1)
    except OSError as e:
        print(f"Could not read {in_path}: {e}")
        sys.exit(1)
    if not isinstance(data, dict):
        print(f"Input JSON must be a JSON object, got {type(data).__name__}: {in_path}")
        sys.exit(1)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    build(data, out_path, chart_dir=in_path.parent / "rendered")
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
