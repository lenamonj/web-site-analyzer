#!/usr/bin/env python3
"""
Offline unit tests for build_exec_report.py.

Builds a report from a synthetic data dict into a temp directory, reopens it
with python-docx, and asserts the document structure: masthead, glance tiles,
callout, chip fills, finding order, footer page field, and the evidence
appendix. Skipped entirely when python-docx is not installed (the scanners
must stay stdlib-only; only this builder depends on docx).

Run from this directory:
    python -m unittest test_exec_report
"""

import contextlib
import io
import json
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

try:
    from docx import Document
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

if HAVE_DOCX:
    import build_exec_report as ber

try:
    import report_charts
    HAVE_MPL = report_charts.HAVE_MPL
except ImportError:
    HAVE_MPL = False

SAMPLE = {
    "site": "example.com",
    "target_url": "https://example.com",
    "date": "2026-07-02",
    "scope": {"pages_reviewed": 12, "method": "Passive external scan"},
    "progress": {"previous_date": "2026-06-01", "new_issues": 1, "resolved_issues": 4},
    "web_vitals": {"source": "field", "captured_note": "Real Chrome users, 28-day p75 (CrUX)",
                   "metrics": [{"label": "LCP", "value": "0.9s", "rating": "Good",
                                "target": "good is 2.5s or less"},
                               {"label": "CLS", "value": "0.31", "rating": "Poor",
                                "target": "good is 0.10 or less"},
                               {"label": "INP", "value": "350ms", "rating": "Needs work"}]},
    "key_dates": {"note": "Public certificate and domain-registration facts, passively measured.",
                  "items": [{"label": "SSL certificate renews", "value": "2026-09-10",
                             "detail": "in 70 days"},
                            {"label": "Domain renews", "value": "2027-04-02",
                             "detail": "in 640 days"},
                            {"label": "Domain registered", "value": "2015-04-02",
                             "detail": "about 11.0 years ago"}]},
    "bottom_line": "The site is sound overall; the one urgent item is consent.",
    "assessment": {
        "strengths": ["TLS and certificates: strong (4 checks pass)"],
        "weaknesses": ["Security posture: poor (6 failing, 1 warnings). Worst: No HSTS."],
    },
    "action_plan": [
        {"priority": "High", "action": "Add an HSTS response header", "affects": "site-wide"},
        {"priority": "Medium", "action": "Right-size page titles", "affects": "2 page(s)"},
    ],
    "scorecard": {
        "overall": "Adequate",
        "rows": [
            {"category": "security", "band": "Poor", "detail": "1/1/6", "score": 0.28},
            {"category": "seo", "band": "Strong", "detail": "9/1/0", "score": 1.0},
        ],
    },
    "findings": [
        {"area": "SEO", "finding": "Low issue", "evidence": "scan.json", "severity": "Low"},
        {"area": "Security", "finding": "High issue", "evidence": "scan.json", "severity": "High"},
        {"area": "Privacy", "finding": "Critical issue", "evidence": "scan.json", "severity": "Critical"},
    ],
    "recommendations": [
        {"rank": 2, "recommendation": "Second", "impact": "B", "effort": "M"},
        {"rank": 1, "recommendation": "First", "impact": "A", "effort": "S"},
    ],
    "quick_wins": ["Fix titles", "Label the form"],
    "evidence": [
        {"caption": "Problem snippet", "code": "<img src=x>", "highlight": "src=x"},
        {"caption": "Missing shot", "image": "planning/_evidence/does_not_exist.png"},
    ],
}


def _cell_fill(cell):
    """The w:fill of a cell's shading, or None."""
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
    if shd is None:
        return None
    return shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")


def _cell_fill_of_run(cell):
    """The w:shd fill on any run inside a cell (the vitals rating chip), or None."""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for r in cell._tc.iter(f"{ns}r"):
        shd = r.find(f"{ns}rPr/{ns}shd")
        if shd is not None:
            return shd.get(f"{ns}fill")
    return None


def _doc_text(document):
    parts = [p.text for p in document.paragraphs]
    for t in document.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


@unittest.skipUnless(HAVE_DOCX, "python-docx not installed")
class TestExecReport(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.TemporaryDirectory()
        cls.out = Path(cls.tmp.name) / "report.docx"
        ber.build(SAMPLE, cls.out)
        cls.doc = Document(str(cls.out))
        cls.text = _doc_text(cls.doc)

    @classmethod
    def tearDownClass(cls):
        cls.tmp.cleanup()

    def test_cover_carries_site_kicker_and_posture(self):
        self.assertIn("example.com", self.text)
        self.assertIn("WEBSITE REVIEW  /  EXECUTIVE REPORT", self.text)
        # The measured overall posture appears on the cover as an uppercase chip.
        self.assertIn("OVERALL POSTURE", self.text)
        self.assertIn("ADEQUATE", self.text)
        # The cover ends with a page break so content starts on its own page.
        body_xml = self.doc.element.body.xml
        self.assertIn('w:br w:type="page"', body_xml)

    def test_cover_contents_list_names_rendered_sections_in_order(self):
        self.assertIn("IN THIS REPORT", self.text)
        idx = self.text.index("IN THIS REPORT")
        listing = self.text[idx:idx + 600]
        names = ("Executive summary", "Measured posture", "Core Web Vitals",
                 "Key dates", "Key findings hurting the site",
                 "Preferred recommendations", "Quick wins", "Evidence appendix")
        positions = []
        for name in names:
            pos = listing.find(name)
            self.assertNotEqual(pos, -1, f"{name!r} missing from the cover contents list")
            positions.append(pos)
        # P63: the contents list must name the sections in rendering ORDER, not just
        # contain them; each name's offset must strictly increase, so a shuffled
        # contents list fails rather than passing on membership alone.
        self.assertEqual(positions, sorted(positions),
                         f"cover contents list out of order: {list(zip(names, positions))}")

    def test_running_header_skips_the_cover(self):
        section = self.doc.sections[0]
        self.assertTrue(section.different_first_page_header_footer)
        self.assertIn("example.com   |   Website Review   |   2026-07-02",
                      section.header.paragraphs[0].text)
        # The cover's own header stays empty.
        self.assertEqual(section.first_page_header.paragraphs[0].text.strip(), "")

    def test_glance_tiles_report_counts_from_data(self):
        tiles = self.doc.tables[0]
        self.assertEqual(len(tiles.columns), 4)
        tile_text = "\n".join(p.text for c in tiles.rows[0].cells for p in c.paragraphs)
        self.assertIn("Adequate", tile_text)          # overall band
        self.assertIn("3", tile_text)                 # findings count
        self.assertIn("1 Critical / 1 High / 1 Low", tile_text)
        # With scope present the fourth tile shows pages reviewed.
        self.assertIn("PAGES REVIEWED", tile_text)
        self.assertIn("12", tile_text)

    def test_scope_line_and_progress_strip(self):
        self.assertIn("12 pages reviewed  |  Passive external scan", self.text)
        self.assertNotIn("page(s)", self.text.replace("2 page(s)", ""))  # data passes through; no builder-made "(s)"
        self.assertIn("Since the previous review (2026-06-01):", self.text)
        self.assertIn("4 resolved", self.text)
        self.assertIn("1 new", self.text)

    def test_cover_counts_measured_bands(self):
        # The cover states the measured band mix, counted from the scorecard rows.
        self.assertIn("2 areas measured: 1 Strong, 1 Poor", self.text)

    def test_section_notes_introduce_tables(self):
        self.assertIn("Each area rolls its automated checks into a posture band."
                      " The bar is the measured pass score, 0 to 1.", self.text)
        self.assertIn("Ordered by severity. The evidence column names every "
                      "affected page; paths are on the reviewed site.", self.text)

    def test_web_vitals_panel_renders_values_and_source(self):
        self.assertIn("CORE WEB VITALS", self.text)  # section headings render uppercase
        self.assertIn("0.9s", self.text)
        self.assertIn("350ms", self.text)
        self.assertIn("Real Chrome users, 28-day p75 (CrUX)", self.text)
        # The published Good threshold renders beside the rating chip when the
        # data supplies one; a metric without a target adds nothing.
        self.assertIn("good is 2.5s or less", self.text)
        self.assertIn("good is 0.10 or less", self.text)
        # The rating chips carry the vitals fill colors.
        panel = next(t for t in self.doc.tables
                     if any("0.9s" in c.text for c in t.rows[0].cells))
        chip_fills = {_cell_fill_of_run(c) for c in panel.rows[0].cells}
        self.assertIn(ber.VITALS_FILL["Good"], chip_fills)
        self.assertIn(ber.VITALS_FILL["Poor"], chip_fills)

    def test_key_dates_panel_renders_cert_and_domain_dates(self):
        self.assertIn("KEY DATES", self.text)  # section heading, uppercased
        self.assertIn("2026-09-10", self.text)   # cert renewal date
        self.assertIn("2027-04-02", self.text)   # domain renewal date
        self.assertIn("about 11.0 years ago", self.text)
        panel = next(t for t in self.doc.tables
                     if any("2026-09-10" in c.text for c in t.rows[0].cells))
        self.assertEqual(len(panel.columns), 3)

    def test_section_headings_are_numbered_and_keep_with_next(self):
        heading = next(p for p in self.doc.paragraphs
                       if p.text.endswith("EXECUTIVE SUMMARY"))
        self.assertTrue(heading.paragraph_format.keep_with_next)
        self.assertTrue(heading.text.startswith("01"))

    def test_scorecard_draws_measured_score_bars(self):
        table = self._table_with_header("MEASURED SCORE")
        bar = table.rows[1].cells[2]  # security, score 0.28
        self.assertEqual(bar.text.count(ber.SCORE_BAR_BLOCK), ber.SCORE_BAR_SEGMENTS)
        self.assertTrue(bar.text.endswith("0.28"))
        # The filled segments carry the band color (Poor), in the first
        # non-empty run (cell.text = "" leaves an empty placeholder run).
        first_run = next(r for r in bar.paragraphs[0].runs if r.text)
        self.assertEqual(str(first_run.font.color.rgb), ber.BAND_FILL["Poor"])
        self.assertEqual(len(first_run.text), round(0.28 * ber.SCORE_BAR_SEGMENTS))

    def test_bottom_line_kicker_renders(self):
        self.assertIn("THE BOTTOM LINE", self.text)

    def test_image_exhibit_renders_in_a_framed_cell(self):
        png_1x1 = (b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00"
                   b"\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc"
                   b"\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82")
        img_path = Path(self.tmp.name) / "shot.png"
        img_path.write_bytes(png_1x1)
        data = dict(SAMPLE)
        data["evidence"] = [{"caption": "Screenshot", "image": str(img_path)}]
        out = Path(self.tmp.name) / "framed.docx"
        ber.build(data, out)
        doc = Document(str(out))
        # The last table is the frame; it must contain exactly one drawing.
        frame_xml = doc.tables[-1]._tbl.xml
        self.assertIn("<pic:pic", frame_xml)

    def test_bottom_line_renders_in_callout(self):
        self.assertIn(SAMPLE["bottom_line"], self.text)

    def test_executive_summary_shows_strengths_and_priorities(self):
        self.assertIn("EXECUTIVE SUMMARY", self.text)
        self.assertIn("STRENGTHS", self.text)
        self.assertIn("PRIORITIES TO FIX", self.text)
        self.assertIn("TLS and certificates: strong", self.text)
        self.assertIn("Security posture: poor", self.text)

    def test_action_plan_renders_when_no_recommendations(self):
        data = {k: v for k, v in SAMPLE.items() if k != "recommendations"}
        out = Path(self.tmp.name) / "plan.docx"
        ber.build(data, out)
        text = _doc_text(Document(str(out)))
        self.assertIn("RECOMMENDED PLAN OF ACTION", text)
        self.assertIn("Add an HSTS response header", text)

    def test_hand_authored_recommendations_win_over_action_plan(self):
        # SAMPLE has recommendations; the auto plan must not also render.
        self.assertIn("PREFERRED RECOMMENDATIONS", self.text)
        self.assertNotIn("RECOMMENDED PLAN OF ACTION", self.text)

    def test_scorecard_band_chips_use_band_fills(self):
        # Chips are compact shaded runs, not full-cell fills: a tall row must
        # show a constant-size badge, never a giant colored slab.
        table = self._table_with_header("POSTURE")
        fills = [_cell_fill_of_run(row.cells[1]) for row in table.rows[1:]]
        self.assertEqual(fills, [ber.BAND_FILL["Poor"], ber.BAND_FILL["Strong"]])
        self.assertTrue(all(_cell_fill(row.cells[1]) is None for row in table.rows[1:]),
                        "band cells must not be cell-filled; the chip is a run")

    def test_findings_sorted_most_severe_first_with_chips(self):
        table = self._table_with_header("SEVERITY")
        sevs = [row.cells[0].text.strip() for row in table.rows[1:]]
        self.assertEqual(sevs, ["Critical", "High", "Low"])
        self.assertEqual(_cell_fill_of_run(table.rows[1].cells[0]),
                         ber.SEVERITY_FILL["Critical"])
        self.assertIsNone(_cell_fill(table.rows[1].cells[0]))

    def test_recommendations_sorted_by_rank(self):
        table = self._table_with_header("RECOMMENDATION")
        recs = [row.cells[1].text.strip() for row in table.rows[1:]]
        self.assertEqual(recs, ["First", "Second"])

    def test_footer_has_page_number_field(self):
        footer = self.doc.sections[0].footer
        xml = footer.paragraphs[0]._p.xml
        self.assertIn("PAGE", xml)
        self.assertIn("example.com Website Review", footer.paragraphs[0].text)

    def test_evidence_appendix_numbers_exhibits(self):
        self.assertIn("Exhibit 1.", self.text)
        self.assertIn("<img src=x>", self.text)
        self.assertIn("[missing image:", self.text)

    def test_minimal_data_builds_without_optional_sections(self):
        out = Path(self.tmp.name) / "minimal.docx"
        ber.build({"site": "bare.example"}, out)
        doc = Document(str(out))
        text = _doc_text(doc)
        self.assertIn("bare.example", text)
        # Only the glance tiles remain (the cover holds no tables); every
        # data-driven table (scorecard, findings, recommendations) is skipped.
        self.assertEqual(len(doc.tables), 1)
        self.assertNotIn("THE BOTTOM LINE", text)

    def _table_with_header(self, header_text):
        """The table whose first-row cells include an exact header label.
        Exact match matters: the glance tiles carry labels like OVERALL
        POSTURE that would satisfy a substring lookup."""
        for table in self.doc.tables:
            if any(c.text.strip() == header_text for c in table.rows[0].cells):
                return table
        self.fail(f"No table with a header cell equal to {header_text!r}")


TREND_2Q = {
    "quarters": ["2026-Q2", "2026-Q3"],
    "series": {"overall_score": [0.72, 0.84]},
    "latest_delta": {
        "prev_quarter": "2026-Q2", "quarter": "2026-Q3",
        "scorecard": [{"category": "security", "prev_band": "Adequate",
                       "band": "Strong", "prev_score": 0.7, "score": 0.9,
                       "direction": "improved"},
                      {"category": "seo", "prev_band": "Strong",
                       "band": "Strong", "prev_score": 0.9, "score": 0.9,
                       "direction": "held"}],
        "new_findings": 1, "resolved_findings": 2,
        "resolved_examples": [
            "Accessibility: 2 links with no discernible text.",
            "SEO: No H1 on the page."],
        "pages_scanned": {"prev": 12, "current": 12},
    },
}

TREND_3Q = {
    "quarters": ["2026-Q1", "2026-Q2", "2026-Q3"],
    "series": {"overall_score": [0.61, 0.72, 0.84],
               "security_score": [0.5, 0.7, 0.9],
               "median_lcp_ms": [3400, 2600, 2100],
               "median_weight_kb": [3100.0, 2500.0, 2400.0],
               "broken_links": [9, 4, 3]},
    "latest_delta": TREND_2Q["latest_delta"],
}


@unittest.skipUnless(HAVE_DOCX, "python-docx not installed")
class TestTrendSection(unittest.TestCase):
    def _data(self, trend):
        data = {k: v for k, v in SAMPLE.items() if k != "evidence"}
        data["slug"] = "example-com"
        data["progress"] = dict(SAMPLE["progress"], trend=trend)
        return data

    def _doc(self, data, tmp):
        out = Path(tmp) / "r.docx"
        ber.build(data, out, chart_dir=Path(tmp) / "charts")
        return Document(str(out))

    def test_trend_table_and_named_resolved_render(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = self._doc(self._data(TREND_2Q), tmp)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("Progress this quarter" in t for t in texts))
        headers = [t.rows[0].cells[0].text.strip() for t in doc.tables]
        self.assertIn("AREA", headers)
        self.assertTrue(any("No H1 on the page." in t for t in texts))
        self.assertTrue(any(
            "2 findings resolved since 2026-Q2; 1 new." in t for t in texts))
        self.assertTrue(any("Pages reviewed: 12 in 2026-Q2, 12 in 2026-Q3"
                            in t for t in texts))

    def test_progress_strip_suppressed_when_trend_renders(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = self._doc(self._data(TREND_2Q), tmp)
        texts = [p.text for p in doc.paragraphs]
        self.assertFalse(any("Since the previous review" in t for t in texts))

    def test_two_quarters_embed_no_charts(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = self._doc(self._data(TREND_2Q), tmp)
        self.assertEqual(len(doc.inline_shapes), 0)

    @unittest.skipUnless(HAVE_MPL, "matplotlib not installed")
    def test_three_quarters_embed_charts(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = self._doc(self._data(TREND_3Q), tmp)
        # overall + categories (one drawable category) + metrics figure
        self.assertEqual(len(doc.inline_shapes), 3)

    def test_three_quarter_trend_without_slug_raises(self):
        """K4: chart file names are prefixed with the slug so two clients'
        trend PNGs can never collide under the shared rendered/ directory;
        a missing slug must fail loudly rather than fall back to "site"."""
        data = self._data(TREND_3Q)
        del data["slug"]
        with tempfile.TemporaryDirectory() as tmp:
            with self.assertRaises(ValueError):
                self._doc(data, tmp)

    def test_two_quarter_trend_without_slug_still_builds(self):
        """No charts render below three quarters, so the missing slug never
        matters and the report must still build."""
        data = self._data(TREND_2Q)
        del data["slug"]
        with tempfile.TemporaryDirectory() as tmp:
            doc = self._doc(data, tmp)
        self.assertEqual(len(doc.inline_shapes), 0)

    def test_cover_contents_list_includes_progress_section(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = self._doc(self._data(TREND_2Q), tmp)
        texts = " ".join(p.text for p in doc.paragraphs)
        self.assertIn("Progress this quarter", texts)


@unittest.skipUnless(HAVE_DOCX, "python-docx not installed")
class TestReportLabel(unittest.TestCase):
    def _doc(self, data):
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "r.docx"
            ber.build(data, out)
            return Document(str(out))

    def test_sample_label_renders_on_cover_and_header(self):
        data = dict(SAMPLE, report_label="SAMPLE REPORT",
                    cover_note="Sample copy: prior-quarter history is illustrative.")
        doc = self._doc(data)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("WEBSITE REVIEW  /  SAMPLE REPORT" in t for t in texts))
        self.assertTrue(any(t.strip() == "SAMPLE REPORT" for t in texts))
        self.assertTrue(any("prior-quarter history is illustrative" in t for t in texts))
        self.assertFalse(any("nothing is estimated" in t for t in texts))
        header = doc.sections[0].header
        self.assertIn("SAMPLE REPORT", header.paragraphs[0].text)

    def test_default_build_keeps_executive_kicker_and_pledge(self):
        doc = self._doc(dict(SAMPLE))
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any("WEBSITE REVIEW  /  EXECUTIVE REPORT" in t for t in texts))
        self.assertTrue(any("nothing is estimated" in t for t in texts))
        self.assertFalse(any(t.strip() == "SAMPLE REPORT" for t in texts))


    def test_recommendations_table_keeps_together(self):
        doc = self._doc(dict(SAMPLE))
        table = next(t for t in doc.tables
                     if len(t.rows[0].cells) > 1
                     and t.rows[0].cells[1].text.strip() == "RECOMMENDATION")
        for row in list(table.rows)[:-1]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    self.assertTrue(para.paragraph_format.keep_with_next)


class TestBuilderDependencies(unittest.TestCase):
    """M3: every third-party import in the builder modules must be declared in
    requirements.txt, so a report never fails on an undeclared dependency (the
    matplotlib gap that made M3). This is the builder analog of the scanner
    charter guard (PLAN section 38); before the fix matplotlib was imported by
    report_charts.py but absent from requirements.txt, and this test would fail."""

    REVIEW = Path(__file__).resolve().parent
    ROOT = REVIEW.parents[2]
    IMPORT_TO_PACKAGE = {"docx": "python-docx"}  # import name -> pip package name

    def _third_party_imports(self, path):
        import ast
        local = {p.stem for p in self.REVIEW.glob("*.py")}
        local |= {p.stem for p in (self.REVIEW / "tools").glob("*.py")}
        mods = set()
        tree = ast.parse(path.read_text(encoding="utf-8"))
        for node in ast.walk(tree):
            if isinstance(node, ast.Import):
                mods |= {a.name.split(".")[0] for a in node.names}
            elif isinstance(node, ast.ImportFrom) and node.level == 0 and node.module:
                mods.add(node.module.split(".")[0])
        return {m for m in mods
                if m not in sys.stdlib_module_names and m not in local}

    def _declared_packages(self):
        pkgs = set()
        for line in (self.ROOT / "requirements.txt").read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            name = line
            for sep in ("<", ">", "=", "!", "~", " "):
                name = name.split(sep)[0]
            pkgs.add(name.strip().lower())
        return pkgs

    def test_builder_third_party_imports_are_declared(self):
        declared = self._declared_packages()
        found = set()
        for module in ("build_exec_report.py", "report_charts.py"):
            found |= self._third_party_imports(self.REVIEW / module)
        self.assertTrue(found, "no third-party imports discovered; the ast/glob "
                               "scan is broken (it should find docx and matplotlib)")
        for imp in found:
            pkg = self.IMPORT_TO_PACKAGE.get(imp, imp).lower()
            self.assertIn(pkg, declared,
                          f"builder imports {imp!r} ({pkg}) but requirements.txt "
                          f"does not declare it")
        self.assertIn("python-docx", declared)
        self.assertIn("matplotlib", declared)


@unittest.skipUnless(HAVE_DOCX, "python-docx not installed")
class TestBuildMainInputGuard(unittest.TestCase):
    """P11: the CLI must reject a structurally wrong input (a top-level JSON
    array) with a clear message and a nonzero exit, not a raw traceback."""

    def _run_main(self, argv):
        orig = sys.argv
        sys.argv = argv
        buf = io.StringIO()
        try:
            with contextlib.redirect_stdout(buf):
                ber.main()
            code = 0
        except SystemExit as e:
            code = e.code
        finally:
            sys.argv = orig
        return code, buf.getvalue()

    def test_top_level_list_exits_nonzero_with_message(self):
        with tempfile.TemporaryDirectory() as td:
            bad = Path(td) / "data.json"
            bad.write_text("[1, 2, 3]", encoding="utf-8")
            code, out = self._run_main(["build_exec_report.py", str(bad), str(Path(td) / "o.docx")])
        self.assertEqual(code, 1)
        self.assertIn("must be a JSON object", out)

    def test_invalid_json_syntax_exits_with_a_clear_message(self):
        # Q9: a syntactically invalid file (a trailing comma is the common slip)
        # must give a clear "Invalid JSON" message and exit 1, not a raw traceback.
        with tempfile.TemporaryDirectory() as td:
            bad = Path(td) / "data.json"
            bad.write_text("{ this is not valid json ", encoding="utf-8")
            code, out = self._run_main(["build_exec_report.py", str(bad), str(Path(td) / "o.docx")])
        self.assertEqual(code, 1)
        self.assertIn("Invalid JSON", out)

    def test_directory_input_exits_with_a_clear_message(self):
        # Q16: pointing the input at a directory must give a clear message and exit
        # 1 (is_file() gate), not a raw PermissionError/OSError traceback.
        with tempfile.TemporaryDirectory() as td:
            code, out = self._run_main(["build_exec_report.py", td, str(Path(td) / "o.docx")])
        self.assertEqual(code, 1)
        self.assertIn("not found", out)

    def test_valid_dict_still_builds(self):
        with tempfile.TemporaryDirectory() as td:
            good = Path(td) / "data.json"
            good.write_text(json.dumps(SAMPLE), encoding="utf-8")
            out_docx = Path(td) / "o.docx"
            code, _ = self._run_main(["build_exec_report.py", str(good), str(out_docx)])
            self.assertEqual(code, 0)
            self.assertTrue(out_docx.is_file())

    def test_mixed_rank_types_sort_numerically_without_crashing(self):
        # P32: string ranks ("1", "10") mixed with an int and a missing rank must
        # not make sorted() compare str and int; they order numerically with a
        # non-numeric or absent rank last.
        data = dict(SAMPLE, recommendations=[
            {"rank": "10", "recommendation": "ten", "impact": "a", "effort": "s"},
            {"rank": 2, "recommendation": "two", "impact": "b", "effort": "m"},
            {"recommendation": "none", "impact": "c", "effort": "l"},
            {"rank": "1", "recommendation": "one", "impact": "d", "effort": "s"},
        ])
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build(data, out)  # must not raise
            doc = Document(str(out))
            table = next(t for t in doc.tables
                         if any(c.text.strip() == "RECOMMENDATION" for c in t.rows[0].cells))
            order = [row.cells[1].text.strip() for row in table.rows[1:]]
        self.assertEqual(order, ["one", "two", "ten", "none"])

    def test_non_string_scalars_do_not_crash_the_build(self):
        # P35: a hand-authoring slip that types site/overall/a quick_win as a number
        # (or a stray dict overall) must stringify, not crash the only deliverable.
        cases = [
            {"site": 123, "target_url": "https://x/", "scorecard": {"overall": "Weak", "rows": []}},
            {"site": "x", "scorecard": {"overall": 0.9, "rows": []}},
            {"site": "x", "scorecard": {"overall": {"band": "Weak"}, "rows": []}},
            {"site": "x", "scorecard": {"overall": "Weak", "rows": []}, "quick_wins": [123, "ok"]},
            {"site": None, "scorecard": {"overall": "Weak", "rows": []}},
            # P38: a key-date label typed as null/number must not crash .upper()
            {"site": "x", "key_dates": {"items": [{"label": None, "value": "2026-01-01"}]}},
            {"site": "x", "key_dates": {"items": [{"label": 123, "value": "2026-01-01"}]}},
            # P39: a non-string evidence code/image must not crash split()/Path()
            {"site": "x", "evidence": [{"caption": "c", "code": 123}]},
            {"site": "x", "evidence": [{"caption": "c", "image": 123}]},
            # Q8: a numeric evidence highlight (bare or in a list) must not crash
            # text.find/iteration in the marked-run renderer.
            {"site": "x", "evidence": [{"caption": "c", "code": "listen 5000", "highlight": [5000]}]},
            {"site": "x", "evidence": [{"caption": "c", "code": "listen 5000", "highlight": 5000}]},
            # P54: a numeric scorecard-row detail must not crash re.sub (it runs
            # only when a numeric score is present on the row), and a numeric
            # report_label must not crash .upper().
            {"site": "x", "scorecard": {"overall": "Weak", "rows": [
                {"category": "SEO", "band": "Weak", "score": 0.5, "detail": 2024}]}},
            {"site": "x", "report_label": 2024, "scorecard": {"overall": "Weak", "rows": []}},
            # Q7: a numeric or list scope.method must not crash the cover-page join.
            {"site": "x", "scope": {"pages_reviewed": 5, "method": 2026}},
            {"site": "x", "scope": {"pages_reviewed": 5, "method": ["a", "b"]}},
        ]
        for data in cases:
            with tempfile.TemporaryDirectory() as td:
                ber.build(data, Path(td) / "o.docx")  # must not raise
                self.assertTrue((Path(td) / "o.docx").is_file())

    def test_string_list_report_fields_render_not_crash(self):
        # Q2: recommendations/findings/action_plan are human-authored on top of the
        # draft (the sibling quick_wins IS a plain string list), so authoring one as
        # a list of strings must render the strings, not crash .get() on a non-dict.
        def build_text(data):
            with tempfile.TemporaryDirectory() as td:
                out = Path(td) / "o.docx"
                ber.build(data, out)
                doc = Document(str(out))
                paras = "\n".join(p.text for p in doc.paragraphs)
                cells = " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
                return paras + " " + cells
        self.assertIn("Add a CSP header",
                      build_text({"site": "x", "recommendations": ["Add a CSP header"]}))
        self.assertIn("Weak TLS config",
                      build_text({"site": "x", "findings": ["Weak TLS config"]}))
        self.assertIn("Enable HSTS",
                      build_text({"site": "x", "action_plan": ["Enable HSTS"]}))
        # a mixed dict + string + number list renders the dict and string, no crash
        mixed = build_text({"site": "x", "findings": [
            {"severity": "High", "area": "tls", "finding": "real dict finding"},
            "bare string finding", 123]})
        self.assertIn("real dict finding", mixed)
        self.assertIn("bare string finding", mixed)
        # Q12: the remaining builder-consumed list fields must also normalize -
        # evidence (authored entirely by hand) plus the nested panel lists.
        self.assertIn("home.png caption",
                      build_text({"site": "x", "evidence": ["home.png caption"]}))
        self.assertIn("seo is weak",
                      build_text({"site": "x", "scorecard": {"overall": "Weak", "rows": ["seo is weak"]}}))
        self.assertIn("LCP 2.1s",
                      build_text({"site": "x", "web_vitals": {"metrics": ["LCP 2.1s"]}}))
        self.assertIn("2027-01-01",
                      build_text({"site": "x", "key_dates": {"items": ["2027-01-01"]}}))

    def test_bare_string_list_field_renders_one_item_not_per_character(self):
        # S3: quick_wins/strengths/weaknesses (and the _as_rows fields) given a bare
        # string must render ONE item, never one bullet per character (a str is
        # iterable). Root fix: _as_str_list and the _as_rows string-field guard.
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "quick_wins": "Add HSTS"}, out)
            bullets = [p.text for p in Document(str(out)).paragraphs
                       if p.style.name == "List Bullet"]
        self.assertEqual(bullets, ["Add HSTS"])          # one bullet, not 8 characters
        # the _as_rows path: a bare-string field is one row, not one row per character
        self.assertEqual(len(ber._as_rows("Homepage lacks a redirect", "finding")), 1)
        # assessment strengths/weaknesses as bare strings build without corruption
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "assessment": {"strengths": "TLS is strong",
                                                    "weaknesses": "No HSTS"}}, out)
            text = " ".join(c.text for t in Document(str(out)).tables
                            for r in t.rows for c in r.cells)
        self.assertIn("TLS is strong", text)
        self.assertIn("No HSTS", text)

    def test_as_rows_handles_a_whole_field_of_any_non_list_type(self):
        # T1: _as_rows must treat a WHOLE field that is not a list as ONE item,
        # whatever its type - the S3 fix closed only the bare-string case. A lone
        # dict (a finding authored without the wrapping list) must render as one row
        # preserving its content, NOT be iterated over its keys (silent drop +
        # fabricated key-name rows); a lone scalar must not crash the build.
        self.assertEqual(ber._as_rows(None, "finding"), [])
        self.assertEqual(ber._as_rows(123, "finding"), [{}])          # scalar: one empty row, no crash
        one = ber._as_rows({"severity": "High", "area": "Security",
                            "finding": "No CSP on homepage", "evidence": "https://x/"}, "finding")
        self.assertEqual(one, [{"severity": "High", "area": "Security",
                                "finding": "No CSP on homepage", "evidence": "https://x/"}])
        # end to end: a single-dict findings field renders the REAL finding, drops nothing
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "findings": {"severity": "High", "area": "Security",
                       "finding": "No CSP on homepage", "evidence": "https://x/"}}, out)
            text = " ".join(c.text for t in Document(str(out)).tables
                            for r in t.rows for c in r.cells)
        self.assertIn("No CSP on homepage", text)
        # a scalar findings field builds instead of crashing
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "findings": 123}, out)
            self.assertTrue(out.is_file())

    def test_non_dict_container_field_skips_section_not_crash(self):
        # S4: a container field given a list or scalar must skip its section cleanly,
        # not AttributeError mid-build (the .get() dereferences were unguarded).
        # U1: scope was the 6th top-level container the coercion loop omitted - an
        # AST enumeration confirms these six are ALL the .get()-dereferenced ones.
        for key in ("scorecard", "web_vitals", "key_dates", "assessment", "progress", "scope"):
            with tempfile.TemporaryDirectory() as td:
                out = Path(td) / "o.docx"
                ber.build({"site": "x", key: [{"a": "b"}]}, out)   # a list, not a dict
                self.assertTrue(out.is_file(), key)

    def test_non_dict_scope_builds_and_valid_scope_still_renders(self):
        # U1: scope read via _scope_text and add_glance_tiles; a prose string or a
        # list of page URLs (a plausible hand-author shape) must skip cleanly, and a
        # valid scope dict must still render its method and pages-reviewed tile.
        for bad in ("Homepage and top nav only", ["/a", "/b"], 5):
            with tempfile.TemporaryDirectory() as td:
                out = Path(td) / "o.docx"
                ber.build({"site": "x", "scope": bad, "bottom_line": "x"}, out)
                self.assertTrue(out.is_file(), repr(bad))
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "scope": {"method": "Automated plus manual",
                                              "pages_reviewed": 7}, "bottom_line": "x"}, out)
            text = " ".join(p.text for p in Document(str(out)).paragraphs)
        self.assertIn("Automated plus manual", text)

    def test_unhashable_lookup_key_field_does_not_crash_the_build(self):
        # U3: band/severity/rating/priority/overall/direction are used as color- and
        # sort-lookup keys; a list/dict value (a hand-author slip) must render with the
        # default instead of TypeError: unhashable. Covers every keyed site including
        # the findings sort key (SEVERITY_ORDER) and the trend direction style.
        hostile = {"site": "x",
                   "scorecard": {"overall": ["W"], "rows": [{"category": "S", "band": {"k": 1}, "score": 0.4}]},
                   "findings": [{"area": "A", "finding": "f", "severity": ["High"]}],
                   "action_plan": [{"priority": {"k": 1}, "action": "a", "affects": "x"}],
                   "web_vitals": {"metrics": [{"label": "L", "value": "2s", "rating": [1]}]}}
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build(hostile, out)                      # every unhashable key at once
            self.assertTrue(out.is_file())
        # regression: normal string severities still sort most-severe first
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "findings": [{"area": "A", "finding": "low1", "severity": "Low"},
                                                 {"area": "B", "finding": "high1", "severity": "High"}]}, out)
            cells = [c.text for tb in Document(str(out)).tables for r in tb.rows for c in r.cells]
        hi = next(i for i, c in enumerate(cells) if "high1" in c)
        lo = next(i for i, c in enumerate(cells) if "low1" in c)
        self.assertLess(hi, lo)

    def test_non_dict_nested_progress_trend_skips_section_not_crash(self):
        # T2: the S4 coercion covered the top-level containers but not the nested
        # progress.trend, which add_trend_section reads with .get(). A non-dict trend
        # (string/list/scalar) must skip the Progress section, not AttributeError.
        for bad in ("oops", [1, 2, 3], 123, True):
            with tempfile.TemporaryDirectory() as td:
                out = Path(td) / "o.docx"
                ber.build({"site": "x", "progress": {"trend": bad, "posture": "improving"}}, out)
                self.assertTrue(out.is_file(), repr(bad))
        # a valid trend dict still renders its Progress this quarter section
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "slug": "x", "progress": {"trend": {
                "latest_delta": {"scorecard": [{"category": "Security", "prev": "Weak",
                                                "current": "Adequate"}],
                                 "resolved_findings": 2, "new_findings": 1,
                                 "prev_quarter": "2026 Q1", "quarter": "2026 Q2"},
                "quarters": ["2026 Q1", "2026 Q2"]}}}, out)
            text = " ".join(p.text for p in Document(str(out)).paragraphs)
        self.assertIn("Progress this quarter", text)

    def test_keyed_object_of_item_dicts_renders_every_item_not_one_bogus_row(self):
        # U4 (structural, via normalize()): findings/recommendations/action_plan
        # authored as a keyed OBJECT of item-dicts ({"f1": {...}, "f2": {...}}) is a
        # keyed collection - normalize() flattens it to its VALUES so EVERY item
        # renders (the no-silent-drop rule), instead of _as_rows iterating the
        # object's keys into one blank Low row and dropping every real finding.
        self.assertEqual(
            ber._as_collection({"f1": {"finding": "No H1"}, "f2": {"finding": "Cert expired"}}),
            [{"finding": "No H1"}, {"finding": "Cert expired"}])
        # a lone item dict has STRING values, not dict values, so it is not a
        # collection and passes through for _as_rows to render as one row (T1).
        self.assertEqual(ber._as_collection({"finding": "No H1", "severity": "High"}),
                         {"finding": "No H1", "severity": "High"})
        data = {"site": "x",
                "findings": {"a": {"area": "SEO", "finding": "No H1 on homepage", "severity": "High"},
                             "b": {"area": "TLS", "finding": "Certificate expired", "severity": "High"}},
                "recommendations": {"r1": {"recommendation": "Add a single H1"},
                                    "r2": {"recommendation": "Renew the certificate"}}}
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build(data, out)
            text = " ".join(c.text for t in Document(str(out)).tables for r in t.rows for c in r.cells)
        for expected in ("No H1 on homepage", "Certificate expired",
                         "Add a single H1", "Renew the certificate"):
            self.assertIn(expected, text)          # nothing dropped
        # a single finding dict (no wrapping list, string values) still renders as
        # ONE row preserving its content - the T1 behavior must not regress.
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "findings": {"area": "Security", "finding": "No CSP on homepage",
                                                 "severity": "High"}}, out)
            text = " ".join(c.text for t in Document(str(out)).tables for r in t.rows for c in r.cells)
        self.assertIn("No CSP on homepage", text)

    def test_mixed_keyed_object_flattens_and_drops_nothing(self):
        # V2: a keyed object whose values are NOT all dicts (item dicts plus a stray
        # comment/annotation key, or a bare-string sibling) is still a keyed
        # collection. _as_collection keys on ANY value being a dict, so the object is
        # flattened to its values - every dict item survives (the hard no-silent-drop
        # rule) and a non-dict entry renders as a visible text row rather than
        # collapsing the whole object into one bogus row that drops the real findings.
        self.assertEqual(
            ber._as_collection({"f1": {"finding": "A"}, "f2": {"finding": "B"}, "note": "c"}),
            [{"finding": "A"}, {"finding": "B"}, "c"])
        # an empty dict is no items -> [] (W1); an all-scalar lone record is still
        # NOT a collection and passes through unchanged
        self.assertEqual(ber._as_collection({}), [])
        self.assertEqual(ber._as_collection({"severity": "High", "finding": "f"}),
                         {"severity": "High", "finding": "f"})
        data = {"site": "x", "findings": {
            "f1": {"area": "SEO", "finding": "Missing title tag", "severity": "High"},
            "f2": {"area": "A11y", "finding": "Low contrast text", "severity": "Low"},
            "note": "double-check f2 before sending"}}
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build(data, out)
            text = " ".join(c.text for t in Document(str(out)).tables for r in t.rows for c in r.cells)
        self.assertIn("Missing title tag", text)                 # real finding kept
        self.assertIn("Low contrast text", text)                 # real finding kept
        self.assertIn("double-check f2 before sending", text)    # comment visible, not silently dropped

    def test_empty_object_list_field_skips_section_like_empty_list(self):
        # W1: a list field authored as an empty OBJECT ({}) must skip its section
        # exactly like [], not render one spurious blank row. _as_collection returns
        # [] for an empty dict (an empty object is no items), so _as_rows yields no
        # rows and the section is omitted.
        self.assertEqual(ber._as_collection({}), [])
        for empty in ({}, []):
            with tempfile.TemporaryDirectory() as td:
                out = Path(td) / "o.docx"
                ber.build({"site": "x", "findings": empty}, out)
                headings = [p.text.lower() for p in Document(str(out)).paragraphs]
            self.assertFalse(any("findings hurting" in h for h in headings), repr(empty))
        # nested empty-object list fields skip their sections too (no blank row)
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "scorecard": {"overall": "Weak", "rows": {}},
                       "web_vitals": {"metrics": {}}, "key_dates": {"items": {}}}, out)
            headings = [p.text.lower() for p in Document(str(out)).paragraphs]
        for absent in ("measured posture", "core web vitals", "key dates"):
            self.assertFalse(any(absent in h for h in headings), absent)
        # a non-empty finding still renders its section (guard against over-skipping)
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "findings": [{"area": "A", "finding": "real one",
                                                  "severity": "High"}]}, out)
            headings = [p.text.lower() for p in Document(str(out)).paragraphs]
        self.assertTrue(any("findings hurting" in h for h in headings))

    def test_partial_trend_dict_with_bad_nested_subfields_builds_clean(self):
        # U5 (structural, via normalize()): add_trend_section / add_trend_table read
        # trend.latest_delta, latest_delta.pages_scanned, latest_delta.scorecard, and
        # trend.quarters; a truthy non-dict/non-list (a partial hand-authored trend)
        # would crash. normalize() coerces every one so the section degrades cleanly.
        for bad_trend in ({"latest_delta": "corrupt", "quarters": "x"},
                          {"latest_delta": {"scorecard": "nope", "pages_scanned": "no"},
                           "quarters": 5},
                          {"latest_delta": [1, 2], "quarters": {"a": 1}}):
            with tempfile.TemporaryDirectory() as td:
                out = Path(td) / "o.docx"
                ber.build({"site": "x", "slug": "x", "progress": {"trend": bad_trend}}, out)
                self.assertTrue(out.is_file(), repr(bad_trend))

    def test_trend_scorecard_with_non_dict_rows_builds_clean(self):
        # V3: U5 coerced the trend scorecard CONTAINER to a list but not its ITEMS, so
        # a non-dict row (str/None) crashed add_trend_table's row.get(). _normalize_trend
        # now routes the scorecard through _as_rows (like the top-level scorecard.rows),
        # so a corrupt row degrades to a dict and the build never crashes.
        for bad_scorecard in (["SEO improved", {"category": "TLS", "prev_band": "Weak", "band": "Strong"}],
                              [None, {"category": "TLS"}],
                              [123, "x"]):
            with tempfile.TemporaryDirectory() as td:
                out = Path(td) / "o.docx"
                ber.build({"site": "x", "slug": "x", "progress": {"trend": {
                    "quarters": ["Q1", "Q2"], "latest_delta": {"scorecard": bad_scorecard}}}}, out)
                self.assertTrue(out.is_file(), repr(bad_scorecard))
        # a valid trend scorecard still renders its QoQ row
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build({"site": "x", "slug": "x", "progress": {"trend": {
                "quarters": ["2026 Q1", "2026 Q2"], "latest_delta": {
                    "scorecard": [{"category": "Security", "prev_band": "Weak", "band": "Adequate"}],
                    "prev_quarter": "2026 Q1", "quarter": "2026 Q2"}}}}, out)
            text = " ".join(c.text for t in Document(str(out)).tables for r in t.rows for c in r.cells)
        self.assertIn("Security", text)

    def test_non_finite_score_renders_not_measured_not_crash(self):
        # S5: a NaN or +/-Infinity score (json.loads accepts both, so one can
        # round-trip in from upstream) must render "not measured" instead of
        # crashing round() in the score bar; a finite sibling still draws its bar.
        for bad in (float("nan"), float("inf"), float("-inf")):
            with tempfile.TemporaryDirectory() as td:
                out = Path(td) / "o.docx"
                ber.build({"site": "x", "scorecard": {"overall": "Weak", "rows": [
                    {"category": "Perf", "band": "Adequate", "detail": "ok", "score": 0.7},
                    {"category": "Security", "band": "Poor", "detail": "m", "score": bad}]}}, out)
                text = " ".join(c.text for t in Document(str(out)).tables
                                for r in t.rows for c in r.cells)
            self.assertIn("not measured", text)

    def test_progress_only_renders_under_an_executive_summary_heading(self):
        # P34: progress without a bottom line, assessment, or trend must render its
        # strip under an Executive summary heading, not orphaned under the tiles.
        data = {"site": "x", "target_url": "https://x/", "date": "2026-07-05",
                "scorecard": {"overall": "Weak", "rows": []},
                "progress": {"previous_date": "2026-04-01", "new_issues": 1,
                             "resolved_issues": 2}}
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "o.docx"
            ber.build(data, out)
            texts = [p.text for p in Document(str(out)).paragraphs]
        self.assertTrue(any(t.strip().endswith("Executive summary") for t in texts),
                        "Executive summary heading missing for progress-only data")
        self.assertTrue(any("Since the previous review" in t for t in texts),
                        "progress strip missing")


if __name__ == "__main__":
    unittest.main()
